from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pyodbc
import pandas as pd
import numpy as np
from chatcompletion import ChatCompletion

class DocxGenerator:
    def __init__(self):
        self.doc = Document()
        self.section = self.doc.sections[0]
        self.header = self.section.header

    def add_logo(self, image_path):
        paragraph = self.header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.5))
        run.add_text('\t')

    def add_title(self, text, font_size=24):
        paragraph = self.header.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(font_size)

    def add_paragraph(self, text, font_size=16):
        paragraph = self.header.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.bold = False
        run.font.size = Pt(font_size)

    def add_separator(self):
        paragraph = self.header.add_paragraph("______________________________________________________________")
        paragraph.style.font.bold = False

    def add_content(self, text, font_size=16):
        paragraph = self.doc.add_paragraph(text)
        paragraph.style.font.bold = True
        paragraph.style.font.size = Pt(font_size)

    def add_table_from_data(self, data):
        table = self.doc.add_table(rows=1, cols=len(data.columns))
        table.style = 'Table Grid'

        # Add column names to the table
        for i, column_name in enumerate(data.columns):
            cell = table.cell(0, i)
            cell.text = column_name
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # Add data rows to the table
        for _, row in data.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cell = cells[i]
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    def generate_docx(self, image_path, title_text, paragraph_text, data):
        self.add_logo(image_path)
        self.add_title(title_text, font_size=24)
        self.add_paragraph(paragraph_text, font_size=16)
        self.add_separator()
        self.add_content("1.ตัวอย่างการจัดทำเอกสารลองดึงข้อมูลจากฐานข้อมูลและสร้างกราฟแสดงผล", font_size=16)
        self.add_table_from_data(data)
        self.add_content("2.ตัวอย่างข้อมูลที่ได้จาก OpenAI API", font_size=16)
        response = ChatCompletion(f'guess what my data is about {data.columns}')
        self.add_content(str(response), font_size=12)
        self.doc.save('output1.docx')

# Usage example
image_path = "D:\AI Automation\AI-Automation\data\KTB-Logo.jpg"
title_text = "บันทึก"
paragraph_text = "เสนอคณะกรรมการขอทำ Project 10X ให้ประเทศไทยไม่เหมือนเดิมอีกต่อไป"
data = pd.DataFrame(np.random.randint(0, 10, size=(5, 3)), columns=['Column 1', 'Column 2', 'Column 3'])

docx_generator = DocxGenerator()
docx_generator.generate_docx(image_path, title_text, paragraph_text, data)
